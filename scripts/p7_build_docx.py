# -*- coding: utf-8 -*-
"""Phase 7: Abgabe-Dokumentation (python-docx).

Aufbau entlang eines durchgehenden roten Fadens: These 'Bildungserfolg wird lokal
entschieden' + Daten-Flow INPUT -> OUTPUT -> UEBERGANG -> ERGEBNIS. Jede Leitfrage
wird als Station auf diesem Fluss beantwortet (Ergebnis, Grafik inline, Interpretation,
Begruendung/Grenzen). Alle Abbildungen sind Original-Screenshots aus Power BI Desktop.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
ROOT=os.path.dirname(os.path.dirname(os.path.abspath(__file__))); CH=os.path.join(ROOT,"charts")
NAVY=RGBColor(0x1E,0x27,0x61); GREY=RGBColor(0x66,0x66,0x66)
d=Document()
st=d.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(11)
def H(t,lvl=1):
    h=d.add_heading(t,lvl)
    for r in h.runs: r.font.color.rgb=NAVY
    return h
def P(t,size=11,bold=False,italic=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic; return p
def PL(label,text,size=11):
    p=d.add_paragraph(); r=p.add_run(label+" "); r.bold=True; r.font.size=Pt(size)
    r2=p.add_run(text); r2.font.size=Pt(size); return p
def code(t):
    p=d.add_paragraph(); r=p.add_run(t); r.font.name="Consolas"; r.font.size=Pt(9); return p
def table(headers,rows):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Light Grid Accent 1"
    for i,hd in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; run=c.paragraphs[0].add_run(hd); run.bold=True; run.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""; run=cells[i].paragraphs[0].add_run(str(v)); run.font.size=Pt(9)
    return t
def pic(fn,wmax=6.3,hmax=4.3,cap=None):
    """Bettet ein Bild ein und skaliert es so, dass weder Breite (wmax) noch Hoehe
    (hmax) in Zoll ueberschritten werden (haelt Hoch- wie Querformat lesbar)."""
    p=os.path.join(CH,fn)
    if not os.path.exists(p):
        P(f"[Abbildung fehlt: {fn}]",9,italic=True); return
    iw,ih=Image.open(p).size; asp=iw/ih
    w=min(wmax,hmax*asp)
    d.add_picture(p,width=Inches(w)); d.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        c=d.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=c.add_run(cap); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=GREY

# ===================== TITEL =====================
ti=d.add_heading("Schulabschluss ist nicht nur Ländersache",0)
for r in ti.runs: r.font.color.rgb=NAVY
P("Data Story zum Self-Service-BI-Projekt",13,italic=True)
P("Max Budde · John Kanto · Aaron Ziegler · HTW Berlin, W2-AA Analytische Anwendungen (Prof. Dr. Kempa)",10)
P("Stand: 02.07.2026 · SSBI-Werkzeug: Power BI Desktop · Datenbasis: offene Daten von Destatis und der Regionalstatistik",10,italic=True)

# ===================== 1. VISION & ROTER FADEN =====================
H("1. Vision und roter Faden")
P("Unsere Leitthese lautet: Bildungserfolg in Deutschland ist kein reines Länderphänomen, sondern wird lokal "
  "entschieden. Der Bundesland-Durchschnitt sagt, in welchem Rahmen sich ein Land bewegt; wo und für wen "
  "Bildung wirklich gelingt oder scheitert, zeigt sich aber erst auf der Kreisebene. Mit offenen Daten von "
  "Destatis und der Regionalstatistik verfolgen wir Schulabschlüsse von der Bundesland- bis auf die Kreisebene "
  "und verknüpfen sie mit Schulstruktur, Bildungsausgaben, Arbeitsmarkt und Einkommen.")
P("Der rote Faden ist ein Daten-Flow in vier Stufen, dem dieser Bericht von Anfang bis Ende folgt:",bold=True)
table(["Stufe","Inhalt","Leitfragen"],[
 ["INPUT","Was fließt hinein: Bildungsausgaben und Schulstruktur","LF5, LF7"],
 ["OUTPUT","Was kommt heraus: Abschlüsse und Abgänge","LF1, LF2, LF3, LF4, LF6"],
 ["ÜBERGANG","Was danach kommt: Berufsschule und Ausbildung","Kontext (berufliche Abgänge)"],
 ["ERGEBNIS","Woran es sich zeigt: Arbeitsmarkt und Einkommen","LF8, LF9"]])
P("Entlang dieses Flusses beantworten wir neun Leitfragen in drei Blöcken. Sie bilden die Kapitelstruktur der "
  "Analyse (Kapitel 4) und kehren im Fazit als Gesamtbild zurück.")
table(["Block","Leitfragen"],[
 ["Der Befund (wo und wie stark?)","LF1 führende Bundesländer 22/23 & 23/24 · LF2 höchster Anteil ohne Hauptschulabschluss · LF3 Streuung der Kreise"],
 ["Die Struktur (wodurch geprägt?)","LF4 Geschlechterunterschied · LF5 Schulartmix · LF6 relativ statt absolut"],
 ["Das Ökonomische (was hängt zusammen?)","LF7 Ausgaben nach Bereich · LF8 Ausgaben und Abschlüsse · LF9 Risiko-Kreise (Bildung, Arbeit, Einkommen)"]])
PL("Leseführung:","Kapitel 2 und 3 legen Datengrundlage und Modell als Fundament. Kapitel 4 ist das Herzstück "
   "und führt Frage für Frage durch den Fluss, jeweils mit Ergebnis, Grafik, Interpretation und Begründung. "
   "Kapitel 5 bis 7 sichern das Ergebnis methodisch ab und ziehen das Fazit. Wer nur die Ergebnisse lesen "
   "möchte, beginnt bei Kapitel 4. Alle Abbildungen sind Original-Ansichten aus dem Power-BI-Bericht.")

# ===================== 2. DATENGRUNDLAGE =====================
H("2. Datengrundlage entlang des Flusses")
P("Wir haben ausschließlich öffentlich zugängliche Quellen verwendet (Datenlizenz Deutschland 2.0 / Destatis). "
  "Jede Quelle lässt sich einer Stufe des Flusses zuordnen; damit steht die Datenbasis von Beginn an im Dienst "
  "der Story.")
table(["Fluss-Stufe","Quelle / Tabelle","Ebene","Jahr"],[
 ["INPUT","Ausgaben je Schüler nach Schulart (21711)","Bundesland","2010–2024 / 2023"],
 ["INPUT","Schulen & Schüler nach Schulart (21111-01-03-4)","Kreis","2023"],
 ["OUTPUT","Abgänge allgemeinbildender Schulen (21111-02-06-4)","DE/BL/RB/Kreis","2023"],
 ["OUTPUT","Abgänge nach Bundesland (Statistischer Bericht 21111-12)","Bundesland","2022 & 2023"],
 ["ÜBERGANG","Berufliche Schulabschlüsse (21121-02-02-4)","Kreis","2023"],
 ["ERGEBNIS","Arbeitslose / Jugend-Arbeitslosenquote (13211-02-05-4)","Kreis","2025"],
 ["ERGEBNIS","Verfügbares Einkommen je Einwohner, VGRdL (82411-01-03-4)","Kreis","2021"],
 ["Hilfsgröße","Bevölkerung nach Altersgruppen (12411-02-03-4)","Kreis","1995–2024"]])
P("Regionalschlüssel (AGS): DG = Deutschland, zweistellig = Bundesland, dreistellig = Regierungsbezirk bzw. "
  "statistische Region, fünfstellig = Kreis. Die Rohdaten liegen als CSV (Regionalstatistik, Windows-1252, "
  "Trennzeichen Semikolon, mehrzeilige Kopfzeilen) sowie als XLSX (Destatis) vor.")
PL("Beispiel und Cross-Source-Prüfung:","Für Schleswig-Holstein ergeben sich 2022/23 = 2.333 und 2023/24 = "
   "2.499 Abgänger ohne Hauptschulabschluss. Derselbe Wert (2.499) folgt unabhängig aus der Regionalstatistik "
   "und aus dem Statistischen Bericht. Solche Querabgleiche haben wir genutzt, um die Aufbereitung abzusichern.")
P("Power Query (M) lädt die offenen Rohdateien direkt (Windows-1252-CSV der Regionalstatistik und die "
  "Destatis-XLSX) und führt sämtliche Aufbereitungsschritte im Modell aus – es gibt keine vorgelagerte "
  "Bereinigung außerhalb des BI-Werkzeugs. Die wichtigsten Transformationen:")
for t in [
 "Encoding Windows-1252 nach UTF-8 (geprüft über eine Mojibake-Gegenprobe).",
 "Sonderzeichen wie '-', '.', 'x' werden als fehlend behandelt, nie als 0.",
 "Umbau von breiten, mehrfach verschachtelten Kopfzeilen in saubere Spalten (Region × Jahr × Merkmal × Wert).",
 "AGS-Normalisierung und Ableitung des Bundesland-Codes aus den ersten zwei Stellen.",
 "Dezimal-Parsing mit Gebietsschema en-US, da die Quell-CSVs den Punkt als Dezimaltrennzeichen nutzen (verhindert den Faktor-10-Fehler bei Quoten, siehe DQ8).",
 "Trimmen leerzeichengefüllter Gebietsnamen für saubere Schlüssel-Joins (DQ10)."]:
    d.add_paragraph(t,style="List Bullet")

# ===================== 3. DATENMODELL =====================
H("3. Das Datenmodell als Rückgrat der Analyse")
P("Damit wir jede Leitfrage auf der jeweils richtigen Ebene beantworten können, steht hinter der Story ein "
  "dimensionales Sternschema (Kimball). Die Präsentations-Vision (S10) nennt vier konzeptionelle Kern-Fakten; "
  "physisch sind es fünf Fakttabellen (die Ausgaben zusätzlich nach Schulart) plus drei Hilfsfakttabellen "
  "(Bevölkerung als Nenner, berufliche Abgänge als Übergang, verfügbares Einkommen für die Ergebnis-Stufe), "
  "also insgesamt 8 Fakttabellen und 4 Dimensionen.")
table(["Typ","Tabelle","Grain / Schlüssel","Fluss-Stufe"],[
 ["Fakt","Abgänge","Region × Jahr × Abschluss × Geschlecht","OUTPUT"],
 ["Fakt","Schule","Region × Jahr × Schulart","INPUT"],
 ["Fakt","Arbeitsmarkt","Region × Jahr","ERGEBNIS"],
 ["Fakt","Ausgaben (gesamt)","Bundesland × Jahr; region_code → *:1","INPUT"],
 ["Fakt","Ausgaben nach Schulart","Bundesland × Schulart × Jahr; region_code → *:1","INPUT"],
 ["Hilf","Bevölkerung","Region × Jahr × Altersgruppe (Nenner)","OUTPUT (relativ)"],
 ["Hilf","Abgänge beruflich","Region × Jahr × Abschluss","ÜBERGANG"],
 ["Hilf","Einkommen (VGRdL)","Region × Jahr (verf. Einkommen je Einwohner, 2021)","ERGEBNIS"],
 ["Dim","Region","region_code (AGS), Ebene, Hierarchie Land→RB→Kreis","alle"],
 ["Dim","Zeit","jahr, Schuljahr, Kalenderjahr","alle"],
 ["Dim","Abschluss","abschluss_key, Rang","OUTPUT"],
 ["Dim","Schulart","Schulart","INPUT"]])
P("Die zentrale, über alle Prozesse konforme Dimension ist die Region: dim_region verbindet über den eindeutigen "
  "region_code alle acht Fakten als reines Sternschema (1:n, Single-Direction). Die Zeit ist aktiv nur an den "
  "Abgängen verknüpft, weil dies die einzige echte Mehrjahres-Analyse ist (Schuljahre 2022/23 und 2023/24); die "
  "übrigen Fakten sind Einzeljahr-Snapshots oder Mehrjahres-Durchschnitte. Den beiden Ausgaben-Tabellen und der "
  "Einkommenstabelle wurde ein region_code ergänzt (Name auf AGS), sodass alle Verknüpfungen als sauberes "
  "region_code → dim_region[region_code] (*:1) über den eindeutigen Schlüssel laufen, ohne Klartext-Namensschlüssel.")
pic("pbi_model.png",6.4,4.9,"Abb. 1: Das Sternschema live in Power BI Desktop – Faktentabellen (fact_abgaenge, "
    "fact_arbeitsmarkt_2025, fact_abgaenge_beruflich, fact_bevoelkerung u. a.) rund um die konforme Dimension "
    "dim_region, verbunden über 1:n-Beziehungen (Modellansicht, Ausschnitt).")
PL("Region-Hierarchie und Gebietsstand:","Damit die Navigation nicht allein an manuellen Ebenen-Filtern hängt, "
   "trägt dim_region eine echte Hierarchie Land → Regierungsbezirk → Kreis (aus dem AGS abgeleitet). Der "
   "in der Vision vorgesehene Drilldown ist damit als Modell-Hierarchie hinterlegt. Weil Regionsnamen nicht "
   "eindeutig sind (neun Dubletten durch Gebietsreformen, siehe DQ9), laufen alle Joins über den eindeutigen "
   "region_code. Als bewusste Entscheidung führen wir den Gebietsstand als Slowly Changing Dimension vom Typ 1 "
   "auf den Stand 2023; eine Typ-2-Historisierung wäre nur bei fortlaufenden Lieferungen mit wechselnden "
   "Gebietsständen sinnvoll und hier Overengineering.")
P("Bus-Matrix (welche Dimension trägt welchen Faktprozess):",bold=True)
table(["Faktprozess","Region","Zeit","Abschluss","Schulart"],[
 ["Abgänge","X","X (aktiv)","X","–"],
 ["Schule","X","(2023)","–","X"],
 ["Arbeitsmarkt","X","(2025)","–","–"],
 ["Ausgaben (gesamt)","X","(Ø/2023)","–","(Attribut)"],
 ["Ausgaben nach Schulart","X","(2023)","–","(Attribut)"],
 ["Bevölkerung (Hilf)","X","(2023)","–","–"],
 ["Abgänge beruflich (Hilf)","X","(2023)","X","–"],
 ["Einkommen (Hilf)","X","(2021)","–","–"]])
P("Region ist die einzige über alle Prozesse konforme Dimension. Zur Additivität: Mengen (anzahl) sind "
  "voll-additiv, die Bevölkerung ist semi-additiv (Bestandsgröße), alle Quoten und Durchschnitte sowie der "
  "Risiko-Score und die Standardabweichung sind nicht-additiv und deshalb als kontextabhängige DAX-Measures "
  "umgesetzt, nicht als vorsummierte Spalten.")
PL("OLAP-Einordnung (warum DAX und nicht MDX):","Power BI Desktop nutzt die Tabular-Engine (VertiPaq), ein "
   "spaltenorientiertes In-Memory-Modell. Konzeptionell entspricht das einem ROLAP-/In-Memory-Ansatz mit "
   "Sternschema, nicht einem vorberechneten MOLAP-Würfel. Abfragesprache ist deshalb DAX (Tabular-nativ); MDX "
   "wäre die Sprache eines multidimensionalen SSAS-Cubes. Die OLAP-Grundoperationen bilden wir dennoch ab: "
   "Slice/Dice über Slicer und Cross-Filter, Roll-up und Drill-down über die Region-Hierarchie, Pivot über "
   "Achsentausch. Für eine Data Story dieser Größe ist Tabular/DAX gegenüber einem Cube die wartungsärmere Wahl.")

# ===================== 4. ANALYSE =====================
H("4. Die Analyse Schritt für Schritt")
P("Wir gehen nun Frage für Frage durch den Fluss. Zu jeder Leitfrage gehört mindestens ein DAX-Measure im "
  "Modell; jedes Ergebnis haben wir zusätzlich unabhängig direkt aus den Rohdaten gegen die amtlichen "
  "Quellwerte nachgerechnet, um Formelfehler auszuschließen. Ein einfaches Beispiel-Measure:")
code('Quote ohne HSA % =\n    DIVIDE ( [Abgänge ohne HSA], [Abgänge] ) * 100')
P("Alle Visuals folgen denselben Gestaltungsregeln: eine Kernbotschaft je Grafik, Mengenachsen ab 0, "
  "barrierearme Okabe-Ito-Farben (über das Report-Theme okabe_ito_theme.json) und Quellenangaben. "
  "Korrelationen werden mit Stichprobengröße n, p-Wert und 95%-Konfidenzintervall angegeben. Die folgenden "
  "Abbildungen sind die zugehörigen Original-Berichtsseiten aus Power BI.")

H("4.1 Der Befund: Wo und wie stark? (OUTPUT)",2)
P("LF1 – Welche Bundesländer führen 2022/23 und 2023/24 bei Abgängen ohne Abschluss?",bold=True)
PL("Ergebnis:","Sachsen-Anhalt führt in beiden Jahren klar. Der Anteil der Schulabgänger ohne "
   "Hauptschulabschluss steigt von 11,28 % (2022/23) auf 12,66 % (2023/24). Dahinter folgen 2023/24 Thüringen "
   "(10,06 %), Schleswig-Holstein (9,94 %), Mecklenburg-Vorpommern (9,93 %) und Bremen (9,83 %); am niedrigsten "
   "liegt Bayern mit rund 5,4 %.")
pic("pbi/pbi_lf1.png",cap="Abb. 2 (LF1): Berichtsseite LF1 – Anteil ohne Hauptschulabschluss je Bundesland: "
    "Balken-Rangliste, Deutschlandkarte auf Bundeslandebene (Blasengröße = Quote) sowie Land- und Stadtstaat-Slicer.")
PL("Interpretation:","Es gibt ein stabiles Muster aus ostdeutschen Flächenländern plus Bremen und "
   "Schleswig-Holstein. Bemerkenswert ist nicht nur die Rangfolge, sondern der Trend: In den meisten führenden "
   "Ländern steigt der Anteil von 2022/23 auf 2023/24, das Problem verschärft sich also leicht. Das Grundniveau "
   "wird erkennbar durch landesspezifische Strukturen geprägt.")
PL("Warum so:","Wir zeigen bewusst beide Schuljahre, weil erst der Vergleich den Trend sichtbar macht; genau "
   "dafür ist die Zeitdimension aktiv an den Abgängen verknüpft. Die Bundeslandebene beantwortet aber nur den "
   "Rahmen, nicht das ganze Bild – das führt direkt zu LF2 und LF3.")

P("LF2 – Wo ist der Anteil ohne Hauptschulabschluss am höchsten?",bold=True)
PL("Ergebnis:","Auf Kreisebene liegen die Spitzenwerte bei rund 15 bis 17 %, also etwa dreimal so hoch wie in "
   "den besten Kreisen. An der Spitze stehen Anhalt-Bitterfeld (16,78 %) und Pirmasens (16,50 %), gefolgt von "
   "Burgenlandkreis (15,12 %), Mansfeld-Südharz (14,88 %), Wittenberg (14,79 %), Dessau-Roßlau (14,78 %), "
   "Halle (14,71 %) und Suhl (14,42 %).")
pic("pbi/pbi_lf2.png",cap="Abb. 3 (LF2): Berichtsseite LF2 – Kreis-Hotspots ohne Hauptschulabschluss: Balken-Rangliste, "
    "interaktive Deutschlandkarte (Blasengröße = Quote ohne HSA je Kreis) sowie Slicer für Bundesland und Ost/West.")
PL("Interpretation:","Das Problem konzentriert sich stark regional; fast alle Hotspots liegen in Sachsen-Anhalt, "
   "dazu Thüringen. Der Ausreißer Pirmasens (kreisfreie Stadt in Rheinland-Pfalz) zeigt aber, dass es kein "
   "reines Ostphänomen ist, sondern überall dort auftritt, wo strukturschwache Städte und Kreise liegen.")
PL("Warum so:","Wir filtern das Visual bewusst auf ebene=KR, damit ausschließlich Kreise verglichen werden; "
   "ohne diesen Filter würden im Modell auch Bundesland- und Bezirkszeilen mitgezählt und die Werte verfälscht.")

P("LF3 – Länder- oder Kreisproblem: Wie stark streuen die Kreise?",bold=True)
PL("Ergebnis:","Es ist beides. Innerhalb der Länder streuen die Kreise erheblich. Die größte Streuung hat "
   "Rheinland-Pfalz mit einer Standardabweichung von 2,84 Prozentpunkten und einer Spannweite von 12,71 "
   "Prozentpunkten über 36 Kreise. Auch Thüringen (2,66) und Sachsen-Anhalt (2,58) streuen stark, während "
   "Nordrhein-Westfalen (1,76) vergleichsweise homogen ist.")
pic("pbi/pbi_lf3.png",cap="Abb. 4 (LF3): Berichtsseite LF3 – Streuung der Kreis-Quote (ohne HSA × Abitur je Kreis).")
PL("Interpretation:","Der Landesdurchschnitt verdeckt große kommunale Unterschiede. Genau das ist der Kern "
   "unserer These: Bildungsrisiko ist sowohl ein Länder- als auch ein Kreisproblem, und wirksame Maßnahmen "
   "müssen lokal ansetzen.")
PL("Warum so:","Die Streuung berechnen wir als eigenes Measure (Stichproben-Standardabweichung STDEVX.S über "
   "die Kreise eines Bundeslandes). So ist die Aussage nicht nur eine Grafik, sondern eine im Modell "
   "reproduzierbare Kennzahl.")

H("4.2 Die Struktur: Wodurch geprägt? (INPUT und Verteilung)",2)
P("LF4 – Schneiden Jungen und Mädchen unterschiedlich ab?",bold=True)
PL("Ergebnis:","Ja, mit einem deutlichen und gegenläufigen Gefälle. Ohne Hauptschulabschluss bleiben Jungen "
   "mit 8,40 % gegenüber Mädchen mit 5,78 % (Jungen rund 45 % häufiger). Beim Abitur ist es umgekehrt: Mädchen "
   "erreichen 37,12 %, Jungen 29,34 % (Mädchen rund 26 % häufiger).")
pic("pbi/pbi_lf4.png",cap="Abb. 5 (LF4): Berichtsseite LF4 – Geschlechtergefälle ohne HSA und beim Abitur (DE 2023).")
PL("Interpretation:","Das Geschlecht wirkt an beiden Enden der Verteilung. Jungen sind am unteren Ende "
   "überrepräsentiert, Mädchen am oberen. Das ist ein struktureller, kein regionaler Effekt und tritt über "
   "alle Länder hinweg auf.")
PL("Warum so:","Wir behandeln Geschlecht als Attribut (degenerate dimension) in der Abgänge-Tabelle, nicht als "
   "eigene Dimension. Das reicht für den Geschlechter-Vergleich und hält das Sternschema schlank.")

P("LF5 – Wie prägt der Schulartmix die Abschlussverteilung?",bold=True)
PL("Ergebnis (Gesamtmix):","Von allen Schülern entfallen 35,2 % auf Grundschulen, 25,9 % auf Gymnasien, 13,1 % "
   "auf integrierte Gesamtschulen, 8,8 % auf Realschulen, 6,3 % auf Schularten mit mehreren Bildungsgängen, "
   "3,9 % auf Förderschulen und nur noch 3,8 % auf Hauptschulen.")
PL("Ergebnis (Fokus: abschlussvergebende Schulen):","Grundschulen vergeben keinen der hier untersuchten "
   "Abschlüsse – sie dominieren den Gesamtmix aber mit gut einem Drittel und verdecken so die eigentliche "
   "Verteilung. Klammert man sie aus (Basis: weiterführende, abschlussvergebende Schulen), dominieren die "
   "Gymnasien mit 40,0 %, gefolgt von integrierten Gesamtschulen (20,2 %), Realschulen (13,5 %), Schularten "
   "mit mehreren Bildungsgängen (9,7 %), Förderschulen (6,0 %) und Hauptschulen (5,9 %). Der Bericht zeigt "
   "beide Sichten nebeneinander (eigenes Measure ohne Grundschulen).")
pic("pbi/pbi_lf5.png",cap="Abb. 6 (LF5): Berichtsseite LF5 – Schulartmix je Schulart (DE 2023): links der "
    "Gesamtmix, rechts die Sicht ohne Grundschulen (nur abschlussvergebende Schulen).")
PL("Interpretation:","Erst die Fokus-Sicht macht die strukturelle Erklärung für die hohen Abiturquoten "
   "sichtbar: Vier von zehn Schülern der weiterführenden Schulen besuchen ein Gymnasium, weitere zwei von "
   "zehn eine integrierte Gesamtschule. Der starke Rückgang der klassischen Hauptschule (nur 5,9 %) zugunsten "
   "integrierter und mehrgliedriger Schulformen verschiebt die Abschlusswege. Der Schulartmix ist damit die "
   "INPUT-Seite, die die Abschlussstruktur prägt.")
PL("Warum so:","Der Nenner des Anteils muss die Sammelkategorie 'Insgesamt' ausschließen, sonst zählt die "
   "Gesamtsumme doppelt und alle Anteile halbieren sich. Für die Fokus-Sicht schließt ein zweites Measure "
   "zusätzlich die Grundschulen aus dem Nenner aus, sodass sich die Anteile der abschlussvergebenden Schulen "
   "wieder zu 100 % addieren; beide Visuals filtern auf die nationale Ebene, um Mehrfachzählung zu vermeiden.")

P("LF6 – Ändert sich die Wertung, wenn man relativ statt absolut zählt?",bold=True)
PL("Ergebnis:","Ja, die Rangfolge kippt vollständig. Absolut führen die bevölkerungsreichen Länder: "
   "Nordrhein-Westfalen (11.835), Baden-Württemberg (6.920) und Bayern (6.474). Bezogen auf je 1.000 der 15- "
   "bis 18-Jährigen führen dagegen Sachsen-Anhalt (41,6), Bremen (34,2) und Thüringen (33,0).")
pic("pbi/pbi_lf6.png",cap="Abb. 7 (LF6): Berichtsseite LF6 – Rangwechsel absolut gegenüber je 1.000 (15-18 Jahre).")
PL("Interpretation:","Absolutzahlen bilden vor allem die Bevölkerungsgröße ab und sind für einen fairen "
   "Vergleich irreführend. Erst die relative Betrachtung zeigt, wo das Problem pro Kopf am größten ist, nämlich "
   "in kleinen ostdeutschen Ländern und Bremen.")
PL("Warum so:","Als Bezugsgröße nutzen wir die Bevölkerung der 15- bis 18-Jährigen aus der Hilfsfakttabelle, "
   "gepinnt auf 2023. Die methodische Kernbotschaft von LF6 ist bewusst gewählt: Die Bezugsgröße entscheidet "
   "über die Aussage – ein Muster, das für die spätere Interpretation von LF8 wichtig ist.")

H("4.3 Das Ökonomische: Was hängt zusammen? (INPUT-Kosten und ERGEBNIS)",2)
P("LF7 – Wie verteilen sich die Bildungsausgaben nach Bereich?",bold=True)
PL("Ergebnis:","Die Ausgaben je Schüler und Jahr steigen mit der Schulart (Deutschland 2023): Grundschule "
   "8.400 €, Realschule 9.700 €, Schularten mit mehreren Bildungsgängen 10.600 €, Gymnasium 10.900 € und "
   "integrierte Gesamtschule 11.600 €.")
pic("pbi/pbi_lf7.png",cap="Abb. 8 (LF7): Berichtsseite LF7 – Bildungsausgaben je Schüler nach Schulart und Bundesland (2023).")
PL("Interpretation:","Die weiterführenden und integrierten Schularten sind je Schüler am teuersten, was mit "
   "kleineren Lerngruppen, Fachräumen und längerer Schulzeit plausibel ist; die Grundschule ist am günstigsten. "
   "LF7 liefert die reine Verteilung der INPUT-Kosten und ist die Grundlage für die anschließende Frage nach "
   "einem Zusammenhang mit dem Ergebnis (LF8).")
PL("Warum so und Grenzen:","Wir zeigen die Ausgaben modellgestützt sowohl nach Schulart als auch nach "
   "Bundesland und pinnen sie über jahrgebackene Measures auf 2023. Zu beachten ist die Quellenlage: Destatis "
   "weist für 2023 auf Bundesebene fünf Schularten aus; Hauptschulen (nur einzelne Länder) sowie Förder- und "
   "Sonderschulen fehlen quellseitig und sind daher nicht enthalten. Diese Lücke legen wir offen, statt sie zu "
   "kaschieren.")

P("LF8 – Mehr Ausgaben je Schüler, bessere Abschlüsse?",bold=True)
PL("Was wir zunächst sahen:","Auf den ersten Blick scheint der Zusammenhang zu bestehen: Über alle 16 "
   "Bundesländer korrelieren höhere Ausgaben je Schüler mit höheren Abiturquoten (r = +0,61, p = 0,012).")
pic("pbi/pbi_lf8.png",cap="Abb. 9 (LF8): Berichtsseite LF8 – Ausgaben je Schüler und Abiturquote je Bundesland "
    "(2023). Die Stadtstaaten sind farblich abgesetzt; sie treiben die scheinbar positive Trendlinie.")
PL("Was wir tatsächlich herausgefunden haben:","Dieser Eindruck ist ein Stadtstaaten-Artefakt. Berlin, Hamburg "
   "und Bremen verbinden struktur- und stadtbedingt hohe Ausgaben mit hohen Abiturquoten und ziehen die "
   "Trendlinie nach oben. Nimmt man die drei Stadtstaaten heraus, dreht der Zusammenhang unter den 13 "
   "Flächenländern ins Negative (r = −0,36) und ist nicht mehr signifikant. Wir haben die Punkte deshalb im "
   "Live-Visual nach Stadtstaat und Flächenland eingefärbt, damit dieser Effekt unmittelbar sichtbar wird.")
PL("Womit wir das belegen – und womit nicht:","Für einen belastbaren Beleg ist die Datenlage schlicht zu dünn. "
   "Ausgaben liegen nur je Bundesland vor, also nur 16 Datenpunkte. Entsprechend breit sind die "
   "Konfidenzintervalle: für alle 16 Länder reicht das 95%-Intervall von +0,17 bis +0,85, für die 13 "
   "Flächenländer von −0,76 bis +0,24 und schließt damit die Null ein. Ein einfacher Zusammenhang nach dem "
   "Muster mehr Geld führt zu mehr Abitur lässt sich aus diesen Daten nicht belegen. Korrelation ist nicht "
   "Kausalität, und der eigentliche Confounder ist hier der Stadtstaat-Status.")
PL("Warum wir es trotzdem zeigen:","Ein ehrlich ausgewiesenes Nicht-Ergebnis ist selbst ein Ergebnis. LF8 "
   "widerlegt eine naheliegende, aber falsche These und schärft den Blick dafür, dass Bildungserfolg weniger an "
   "der reinen Ausgabenhöhe hängt als an den lokalen Strukturen – die Brücke zu LF9.")

P("LF9 – Welche Kreise verbinden Bildungsrisiko, Arbeitslosigkeit und niedriges Einkommen?",bold=True)
PL("Ergebnis:","Wir fassen alle drei Dimensionen der Ergebnis-Stufe in einem z-standardisierten Risiko-Score "
   "zusammen: Bildungsrisiko (Quote ohne Hauptschulabschluss 2023), Jugendarbeitslosigkeit (15 bis 25 Jahre, "
   "2025) und niedriges verfügbares Einkommen je Einwohner (VGRdL 2021). Die höchsten Werte erreichen "
   "Gelsenkirchen (8,1), Pirmasens (7,4), Mansfeld-Südharz (6,6), Stendal (6,0) und Bremerhaven (6,0).")
pic("pbi/pbi_lf9.png",cap="Abb. 10 (LF9): Berichtsseite LF9 – Bildungsrisiko × Jugendarbeitslosigkeit je Kreis "
    "und dreidimensionales Risiko-Score-Ranking (inkl. Einkommen); Spitze: Gelsenkirchen (rund 8,1). Interaktiv "
    "filterbar über einen Bundesland-Slicer und einen Einkommens-Schieberegler.")
PL("Interpretation:","Die drei Merkmale ballen sich in denselben Kreisen: strukturschwache Städte und Kreise im "
   "Ruhrgebiet, in Rheinland-Pfalz, Sachsen-Anhalt und an der Unterweser. Niedriges Einkommen korreliert "
   "erwartungsgemäß mit hohem Bildungsrisiko (r = −0,49) und hoher Jugendarbeitslosigkeit (r = −0,59); die "
   "Merkmale zeigen also in dieselbe Richtung. Mit der Einkommensdimension rückt Gelsenkirchen (verfügbares "
   "Einkommen nur rund 17.900 € je Einwohner) an die Spitze, und Bremerhaven kommt neu in die Spitzengruppe. "
   "Diese Kreise sind konkrete Kandidaten für gezielte Förderung.")
PL("Warum wir es so gemacht haben:","Um drei sehr unterschiedliche Größen (Prozent, Prozent, Euro) vergleichbar "
   "zu machen, standardisieren wir jede über die 398 Kreise mit allen drei Kennzahlen (z-Wert je Kreis). Das "
   "Einkommen geht invertiert ein (Mittel minus Kreiswert), weil niedriges Einkommen ein höheres Risiko "
   "bedeutet. Der Score gewichtet die drei Dimensionen gleich. Wir haben die Rangliste gegen andere Gewichtungen "
   "geprüft (gleich, bildungs-, arbeitslosen- und einkommenslastig): Das Führungsduo Gelsenkirchen und Pirmasens "
   "bleibt in allen geprüften Varianten durchgängig in den Top-3 (beide wechseln sich je nach Gewichtung als "
   "Erstplatzierter ab); die weiteren Ränge variieren dagegen (etwa Mansfeld-Südharz, Bremerhaven oder Uckermark). "
   "Damit beantwortet LF9 die Leitfrage vollständig, inklusive der Einkommensdimension, die die Vision als "
   "Ergebnis-Stufe von Anfang an vorsah.")
PL("Grenzen, ehrlich benannt:","Die Datenstände unterscheiden sich (ohne HSA 2023, Jugend-ALQ 2025, Einkommen "
   "2021, da dies die aktuellste VGRdL-Veröffentlichung ist). Der Score ist deshalb als Strukturindikator zu "
   "lesen, nicht als tagesaktuelle Momentaufnahme. Der Kern-Zusammenhang Bildungsrisiko und "
   "Jugendarbeitslosigkeit ist über alle 398 Kreise hoch signifikant (r = +0,58, p < 0,001, 95%-Intervall "
   "+0,51 bis +0,64) und – anders als bei LF8 mit nur 16 Ländern – durch die große Kreiszahl eng geschätzt. "
   "Auch hier gilt jedoch: ein statistischer Zusammenhang ist kein Kausalnachweis.")

# ===================== 5. DATENQUALITÄT =====================
H("5. Datenqualität und Verlässlichkeit")
P("Weil die Aussagen nur so gut sind wie die Daten, haben wir die Qualität systematisch geprüft und die "
  "wichtigsten Befunde behoben.")
table(["Dimension","Methode","Ergebnis"],[
 ["Vollständigkeit","Null-/Missing-Rate je Tabelle","Ausgaben 0 %; Arbeitsmarkt/Bev. 15-16 %; Schule strukturell"],
 ["Plausibilität","Σ Abschlussarten = Insgesamt","Deutschland exakt; 48/53 Regionen exakt (Rest ±5 Rundung)"],
 ["Konsistenz","Σ Kreise = Bundesland","12/14 Länder exakt; SH/NRW ±≤15 (Rundung)"],
 ["Regionale Stabilität","AGS-Schema, Gebietsstand 2023 (SCD Typ 1)","einheitlich; keine relevante Gebietsreform 2023-25"],
 ["Encoding","Erkennung + Konvertierung","Windows-1252 → UTF-8"],
 ["Geheimhaltung","'-','.', 'x' als Missing","umgesetzt, nicht als 0 interpretiert"],
 ["DQ8 Dezimal-Locale","Punkt-Dezimal vs. Modell-Kultur de-DE","behoben: en-US-Typcast (Jugend-ALQ war ×10), == Quellwert"],
 ["DQ9 Mehrdeutige Kreisnamen","9 doppelte region-Namen (523 Codes/514 Namen)","behoben: Join über eindeutigen region_code → *:1 (kein m:n)"],
 ["DQ10 Whitespace-Join","Ausgaben-Gebietsnamen leerzeichengefüllt","behoben: getrimmt → 16/16 Bundesländer matchen"],
 ["DQ11 Jahresbezug","Measures ohne Jahresfilter (Pooling)","behoben: Bezugsjahr 2023 (Bericht-Filter + Measure)"]])
PL("Doppelte Absicherung:","Jede Kennzahl haben wir zweifach ermittelt: einmal als DAX-Measure im Modell und "
   "einmal unabhängig direkt aus den Rohdaten, abgeglichen gegen die amtlichen Quellwerte. Beim Abgleich sind "
   "uns vier Datenfehler aufgefallen (DQ8 bis DQ11), die wir behoben haben; danach stimmten Modell und "
   "Nachrechnung überein. Eine automatische Prüfsuite kontrolliert die Kennzahlen und die Modell-Konsistenz "
   "laufend gegen die dokumentierten Referenzwerte.")
PL("Methodischer Grundsatz:","Korrelation ≠ Kausalität. Sowohl bei LF8 als auch bei LF9 weisen wir "
   "Zusammenhänge aus, ziehen aber keine ursächlichen Schlüsse; plausibler ist jeweils ein Confounder "
   "(Stadtstaat-Status bzw. regionale Wirtschaftsstruktur). Alle Korrelationsaussagen sind daher als Hinweise "
   "auf Strukturzusammenhänge zu lesen, nicht als Kausalnachweise.")
PL("Aggregatebene (ökologischer Fehlschluss):","Alle Auswertungen liegen auf Gebietsebene (Bundesland bzw. "
   "Kreis) vor. Ein Zusammenhang zwischen Regionen darf nicht auf Einzelpersonen übertragen werden – dass in "
   "einkommensschwächeren Kreisen mehr Jugendliche ohne Abschluss bleiben, ist eine Aussage über Kreise, nicht "
   "über einzelne Schülerinnen und Schüler (ökologischer Fehlschluss). Der Risiko-Score (LF9) ist bewusst als "
   "regionaler Strukturindikator formuliert, nicht als individuelle Prognose. Aggregatzusammenhänge können sich "
   "auf Individualebene abschwächen oder umkehren (Simpson-Paradoxon) – genau das zeigt LF8 mit den Stadtstaaten.")

# ===================== 6. WERKZEUG & LERNINHALTE =====================
H("6. Werkzeug-Reflexion und Lerninhalte")
P("Als Self-Service-BI-Werkzeug haben wir Power BI Desktop eingesetzt und dabei die geforderten Lerninhalte "
  "adressiert.")
table(["Kriterium","Bewertung"],[
 ["Datenbeschaffung","Power Query importiert CSV/XLSX direkt (Parameter DataFolder); offene Direktdownloads. Grenze: der GENESIS-Datenbankabruf ist login-pflichtig und wurde über offene Berichte umgangen."],
 ["Modellierung","Sternschema in TMDL (Power BI Project) als Code versioniert; Beziehungen, Hierarchie, Measures. Stärke: dimensionales Modell nativ, TMDL textbasiert und diffbar."],
 ["Abfragesprache","DAX für Measures (Quoten, Anteile, Streuung, Korrelation, 3-dim Risiko-Score); Power Query (M) für Transformationen. Grenze: Kontextübergänge sind fehleranfällig, deshalb gegen die amtlichen Quellwerte validiert."],
 ["Visualisierung","Interaktiver Bericht mit 9 Seiten (LF1–LF9), je mit Kernbotschaft und Erkenntnis-Textbox; Drilldown über die Region-Hierarchie; barrierearmes Okabe-Ito-Theme."]])
P("LI1 Berichte und Berichtsgeneratoren:",bold=True); P("Interaktiver Power-BI-Bericht mit einer Seite je Leitfrage. "
  "Für Interaktivität sorgen zwei geografische Deutschlandkarten (Bundeslandebene auf LF1, Kreisebene auf LF2; "
  "Blasengröße = Quote ohne HSA), auf jeder Seite Slicer (Bundesland, Ost/West, Stadt/Landkreis bzw. "
  "Stadtstaat/Flächenland), ein Einkommens-Schieberegler (LF9, Between-Modus) sowie Drilldown über die "
  "Region-Hierarchie (Land → Regierungsbezirk → Kreis). Alle Visuals cross-filtern sich gegenseitig; auf LF8 lässt "
  "sich der Stadtstaaten-Confounder per Slicer live entfernen, worauf die Korrelation ins Negative kippt.")
P("Hinweis zu den Karten: Es handelt sich um Bing-gestützte Kartenvisuals. Sie rendern nur, wenn in Power BI unter "
  "Optionen → Sicherheit die Option für Kartenvisuals aktiviert ist; dabei werden ausschließlich die öffentlichen "
  "Gebiets-/Kreisnamen zur Geokodierung an Microsoft/Bing übermittelt (keine personenbezogenen Daten). Die übrige "
  "Interaktivität (Slicer, Schieberegler) funktioniert unabhängig davon.",10,italic=True)
P("LI2 Gestaltungsregeln:",bold=True); P("Je Visual eine Kernbotschaft, Mengenachsen ab 0, barrierearme Okabe-Ito-Farben (Kontrast mindestens 4,5:1), Quellenangabe; als prüfbare Akzeptanzkriterien dokumentiert.")
P("LI3 Architekturkonzepte und Modellierungsmuster:",bold=True); P("Dimensionale Modellierung (Kimball-Sternschema) mit klarer Trennung von Fakten und Dimensionen, dokumentiertem Grain, konformer Region-Dimension mit echter Hierarchie, SCD Typ 1, Bus-Matrix und Additivitätsklassifikation (siehe Kapitel 3).")
P("LI4 Data-Vault, bewusste Einordnung:",bold=True)
P("Data Vault (Hubs, Links, Satellites) ist für hochfrequent integrierende, historisierende Enterprise-DWHs mit vielen Quellen und auditierbarer Historie gedacht. Für diese Data Story mit wenigen, statischen, gut strukturierten Stichjahres-Quellen wäre ein Vault Overengineering und würde Join-Komplexität und Abfrage-Performance ohne Mehrwert verschlechtern. Das dimensionale Sternschema ist hier das passende, abfrageoptimierte Muster.")
P("LI5 Self-Service Business Intelligence:",bold=True); P("Power BI Desktop erlaubt es Fachanwenderinnen und Fachanwendern, ohne Programmierung Daten zu laden, zu modellieren, mit DAX auszuwerten und interaktiv zu visualisieren. Die gesamte Pipeline ist self-service-fähig und reproduzierbar.")
P("Live in Power BI (Beispiel LF8, Punkte nach Stadtstaat und Flächenland farblich getrennt):",bold=True)
pic("pbi_report_lf8.png",6.4,4.6,"Abb. 11: Interaktive Berichtsseite LF8 in Power BI Desktop. Die Stadtstaaten (Berlin, "
    "Hamburg, Bremen) sind abgesetzt; sie treiben die scheinbar positive Korrelation. Der Bericht umfasst neun "
    "Seiten LF1 bis LF9.")

# ===================== 7. FAZIT =====================
H("7. Fazit: Bildungserfolg wird lokal entschieden")
P("Am Ende des Flusses lässt sich die Leitthese bestätigen. Auf der OUTPUT-Seite zeigt der Befund ein stabiles "
  "Ländermuster (LF1), aber die eigentliche Musik spielt lokal: Die Kreise streuen innerhalb der Länder stark "
  "(LF3), einzelne Kreise erreichen dreifach so hohe Quoten ohne Abschluss wie andere (LF2), und die relative "
  "Betrachtung dreht die Rangfolge komplett (LF6). Die Struktur dahinter ist erklärbar: der Schulartmix prägt "
  "die Abschlüsse (LF5), und das Geschlechtergefälle wirkt an beiden Enden (LF4).")
P("Auf der ökonomischen Seite können wir die naheliegende These, dass mehr Geld zu besseren Abschlüssen führt, "
  "gerade nicht belegen: Die scheinbar positive Korrelation (r = +0,61) ist ein Stadtstaaten-Artefakt und "
  "verschwindet unter den Flächenländern (LF8); bei nur 16 Bundesländern ist keine dieser Ausgaben-Korrelationen "
  "statistisch belastbar. Wo Bildungsrisiko wirklich zusammenkommt, zeigt erst die Ergebnis-Stufe: Bildungsrisiko, "
  "Jugendarbeitslosigkeit und niedriges Einkommen ballen sich in denselben Kreisen (LF9), allen voran Gelsenkirchen, "
  "Pirmasens und Mansfeld-Südharz. Der Zusammenhang zwischen Bildungsrisiko und Jugendarbeitslosigkeit ist über "
  "alle 398 Kreise hoch signifikant (r = +0,58, p < 0,001).")
P("Bildungserfolg ist damit weniger eine Frage der reinen Ausgabenhöhe als der lokalen Strukturen: Schulartmix, "
  "regionale Hotspots und der Übergang in Arbeitsmarkt und Einkommen. Der Risiko-Score macht diese Kreise "
  "sichtbar und liefert einen konkreten Ansatzpunkt für gezielte Förderung. Alle Aussagen verstehen wir dabei "
  "als Hinweise auf Strukturzusammenhänge, nicht als Kausalnachweise – Korrelation ≠ Kausalität bleibt der "
  "methodische Vorbehalt, der diese Arbeit von der ersten bis zur letzten Seite begleitet.")

# ===================== ANHANG =====================
H("Anhang: Reproduzierbarkeit und Quellen")
P("Alle Ergebnisse sind aus den offenen Quelldaten und dem Power-BI-Projekt (TMDL/PBIP) reproduzierbar. "
  "Begleitende Dokumente: Quellen-Log mit URLs und Abrufdatum (datenquellen_log.md), dimensionales Schema "
  "(dimensionales_schema.md), Analyseabfragen mit den DAX-Measures (analyseabfragen.md), Datenqualitätsbericht "
  "(dq_report.md) sowie die Anforderungs- und Nachweismatrix (traceability.csv).")
P("Direktdownloads (Auszug): regionalstatistik.de/genesisws/downloader/00/tables/{21111-02-06-4, 21111-01-03-4, "
  "21121-02-02-4, 13211-02-05-4, 12411-02-03-4, 82411-01-03-4}_00.csv · destatis.de Statistischer Bericht "
  "Allgemeinbildende Schulen · Ausgaben je Schüler (5217109247005). Abruf 2026-06-29 bis 2026-07-01.",10,italic=True)

out=os.path.join(ROOT,"Schulabschluss_DataStory_Dokumentation.docx")
d.save(out); print("gespeichert:",out)
